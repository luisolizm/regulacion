"""Export utilities: Word checklist, Plotly PNG, Excel Gantt."""
import io, os, tempfile
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ── Word export ───────────────────────────────────────────────────────────────
def _hex_to_rgb(h: str):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def _set_cell_bg(cell, hex_color: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), hex_color.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shading)

def export_checklist_word(data: dict) -> bytes:
    doc = DocxDocument()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    h = doc.add_heading('', 0)
    run = h.add_run(f"LUXEM ENERGÍA — CHECKLIST MAESTRO")
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    run.font.size = Pt(16)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run(data['nombre'])
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

    # Resumen ejecutivo
    doc.add_paragraph()
    p = doc.add_paragraph()
    run3 = p.add_run('Plazo estimado: ')
    run3.bold = True
    p.add_run(f"{data['plazo_min_semanas']}–{data['plazo_max_semanas']} semanas")
    p.add_run('  |  ')
    run4 = p.add_run('Modalidad: ')
    run4.bold = True
    p.add_run(data['modalidad'])
    p.add_run('  |  ')
    run5 = p.add_run('Energético: ')
    run5.bold = True
    p.add_run(data['energia'])

    p2 = doc.add_paragraph()
    run6 = p2.add_run('Ruta crítica: ')
    run6.bold = True
    run6.font.color.rgb = RGBColor(0xCA, 0x6F, 0x1E)
    p2.add_run(data['ruta_critica'])

    # Marco normativo
    doc.add_heading('Marco Normativo Aplicable', 2)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for cell, txt in zip(hdr, ['Instrumento', 'DOF', 'Relevancia']):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(cell, '1B4F72')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, nm in enumerate(data.get('marco_normativo', [])):
        row = t.add_row().cells
        row[0].text = nm['instrumento']
        row[1].text = nm['dof']
        row[2].text = nm['relevancia']
        bg = 'EBF5FB' if i % 2 == 0 else 'FFFFFF'
        for cell in row:
            _set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    # Checklist sections
    doc.add_heading('Checklist de Trámites y Actividades', 1)
    for sec in data['checklist']:
        doc.add_heading(sec['seccion'], 2)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'Table Grid'
        tbl.autofit = False
        # Set column widths
        widths = [Cm(1.4), Cm(4.5), Cm(6.5), Cm(3.2), Cm(3.2)]
        for j, (cell, txt) in enumerate(zip(tbl.rows[0].cells,
                ['Código', 'Ítem', 'Descripción', 'Responsable', 'Referencia Normativa'])):
            cell.text = txt
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(cell, '1B4F72')
            cell.width = widths[j]
        for i, item in enumerate(sec['items']):
            row = tbl.add_row().cells
            row[0].text = item['codigo']
            row[1].text = item['item']
            row[2].text = item['descripcion']
            # Add nota in italic below description
            if item.get('nota') and item['nota'] != '—':
                p = row[2].add_paragraph()
                r = p.add_run(f"📝 {item['nota']}")
                r.italic = True
                r.font.size = Pt(7)
                r.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
            row[3].text = item['responsable']
            row[4].text = item['referencia']
            bg = 'F4F6F7' if i % 2 == 0 else 'FFFFFF'
            for cell in row:
                _set_cell_bg(cell, bg)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
        doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_para.add_run('LUXEM ENERGÍA  |  Documento Interno de Referencia  |  Junio 2026')
    fr.font.size = Pt(7)
    fr.font.color.rgb = RGBColor(0xBC, 0xD3, 0xD8)
    fr.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Plotly PNG export ─────────────────────────────────────────────────────────
def export_gantt_png(fig) -> bytes:
    """Export Plotly figure as PNG bytes using kaleido."""
    try:
        return fig.to_image(format='png', width=1800, height=max(600, fig.layout.height or 800), scale=2)
    except Exception as e:
        # Fallback: return SVG if kaleido not available
        return fig.to_image(format='svg', width=1800)


# ── Excel Gantt export ────────────────────────────────────────────────────────
def export_gantt_excel(data: dict) -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Gantt"

    TOTAL_WEEKS = max(data['plazo_max_semanas'] + 2, 26)
    NAVY = "1B4F72"; WHITE = "FFFFFF"; LGRAY = "F4F6F7"

    def solid(h): return PatternFill("solid", fgColor=h.lstrip('#'))
    def bdr(c="CCCCCC", sz=1): return Side(style="thin", color=c)
    def allB(): b=bdr(); return Border(left=b, right=b, top=b, bottom=b)

    ws.column_dimensions['A'].width = 16
    ws.column_dimensions['B'].width = 32
    ws.column_dimensions['C'].width = 16
    for i in range(TOTAL_WEEKS):
        ws.column_dimensions[get_column_letter(4+i)].width = 2.8
    notes_col = 4 + TOTAL_WEEKS
    ws.column_dimensions[get_column_letter(notes_col)].width = 28

    # Title
    ws.row_dimensions[1].height = 24
    ws.merge_cells(f"A1:{get_column_letter(notes_col)}1")
    c = ws["A1"]
    c.value = f"LUXEM ENERGÍA — GANTT: {data['nombre']}"
    c.font = Font(name="Arial", bold=True, color=WHITE, size=13)
    c.fill = solid(NAVY)
    c.alignment = Alignment(horizontal="left", vertical="center")

    ws.row_dimensions[2].height = 16
    ws.merge_cells(f"A2:{get_column_letter(notes_col)}2")
    c = ws["A2"]
    c.value = (f"Plazo estimado: {data['plazo_min_semanas']}–{data['plazo_max_semanas']} sem  |  "
               f"Modalidad: {data['modalidad']}  |  Ruta crítica: {data['ruta_critica'][:120]}...")
    c.font = Font(name="Arial", color=WHITE, size=8, italic=True)
    c.fill = solid("2E4057")
    c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # Week headers
    HDR = 3
    ws.row_dimensions[HDR].height = 22
    for col, lbl in [(1,"FASE"),(2,"ACTIVIDAD"),(3,"RESPONSABLE")]:
        c = ws.cell(row=HDR, column=col)
        c.value = lbl; c.font = Font(name="Arial", bold=True, color=WHITE, size=9)
        c.fill = solid(NAVY); c.border = allB()
        c.alignment = Alignment(horizontal="center", vertical="center")
    for w in range(1, TOTAL_WEEKS+1):
        col = 3 + w
        c = ws.cell(row=HDR, column=col)
        c.value = str(w); c.font = Font(name="Arial", bold=(w%4==0), color=WHITE, size=8)
        c.fill = solid("243D5C" if w%4==0 else NAVY)
        c.border = allB(); c.alignment = Alignment(horizontal="center", vertical="center")
    c = ws.cell(row=HDR, column=notes_col)
    c.value = "NOTAS / REFERENCIA"; c.font = Font(name="Arial", bold=True, color=WHITE, size=9)
    c.fill = solid(NAVY); c.border = allB()
    c.alignment = Alignment(horizontal="left", vertical="center")

    row = 4
    for fase in data['fases']:
        # Phase header
        ws.row_dimensions[row].height = 18
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=notes_col)
        c = ws.cell(row=row, column=1)
        c.value = f"  {fase['id']}  ·  {fase['nombre']}"
        c.font = Font(name="Arial", bold=True, color=WHITE, size=10)
        c.fill = solid(fase['color'])
        c.alignment = Alignment(horizontal="left", vertical="center")
        c.border = allB()
        row += 1

        for i, act in enumerate(fase['actividades']):
            ws.row_dimensions[row].height = 16
            alt = LGRAY if i % 2 == 0 else WHITE
            for col, val in [(1, ""), (2, act['nombre']), (3, act.get('responsable',''))]:
                c = ws.cell(row=row, column=col)
                c.value = val; c.fill = solid(alt)
                c.font = Font(name="Arial", size=8)
                c.alignment = Alignment(horizontal="left" if col==2 else "center", vertical="center", wrap_text=True)
                c.border = allB()
            s, d = act['inicio'], act['duracion']
            for w in range(1, TOTAL_WEEKS+1):
                col = 3 + w
                c = ws.cell(row=row, column=col)
                c.border = allB()
                if s > 0 and d > 0 and w >= s and w < s + d:
                    c.fill = solid(fase['color'])
                else:
                    c.fill = solid(alt)
            c = ws.cell(row=row, column=notes_col)
            c.value = f"{act.get('nota','')} | {act.get('referencia','')}"
            c.font = Font(name="Arial", size=7, italic=True)
            c.fill = solid(alt); c.border = allB()
            c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            row += 1

    ws.freeze_panes = "D4"
    ws.page_setup.orientation = "landscape"
    ws.page_setup.paperSize = 9

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
